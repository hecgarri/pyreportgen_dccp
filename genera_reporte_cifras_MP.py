# -*- coding: utf-8 -*-
"""
Titulo: Automatizacion reportes regionales
Created on Thu Aug 22 11:58:39 2019

Proceso que automatiza reportes regionales, generando un documento word por cada region.

Obs: Se deben revisar los inputs antes de correr

@author: hugo.gallardo
"""

# =============================================================================
# Extracción en python de datos de comunicado de compras a nivel regional
# =============================================================================

#Seteo de librerias

import pandas as pd
import numpy as np
import sqlalchemy as sa #Para conexión a BD, requerido para usar pd.read_sql()
import urllib #Para formatear string de conexión
import docx
import matplotlib.pyplot as plt 
import matplotlib.patches as mpatches
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
import seaborn as sns
import os
import sys
from pylab import savefig
import itertools

import pyodbc ### another engine to make DB connection and run the queries. Alternative to sqlalchemy
from itertools import repeat


#Transforma codigo a word
#Agregar github


#Conexion a DW

#param = urllib.parse.quote_plus("DRIVER={SQL Server};SERVER=10.34.71.202;UID=datawarehouse;PWD=datawarehouse;DATABASE=DM_Transaccional")#;Trusted_Connection=yes
#conn = sa.create_engine("mssql+pyodbc:///?odbc_connect=%s" % param)

#param = urllib.parse.quote_plus("DRIVER={SQL Server};SERVER=10.34.71.227;UID=eymetric;PWD=Eym3tr1c.2022;DATABASE=DM_Transaccional_2022")#;Trusted_Connection=yes
#conn = sa.create_engine("mssql+pyodbc:///?odbc_connect=%s" % param)

#A DCCPProcurement
#paramAQ = urllib.parse.quote_plus("DRIVER={ODBC Driver 13 for SQL Server};DATABASE=DCCPProcurement;SERVER=10.34.71.145;UID=datawarehouse;PWD=datawarehouse")
#bd = sa.create_engine("mssql+pyodbc:///?odbc_connect=%s" % paramAQ)
#paramAQ = urllib.parse.quote_plus("DRIVER={ODBC Driver 13 for SQL Server};SERVER=10.34.71.146;UID=datawarehouse;PWD=datawarehouse")
#bd = sa.create_engine("mssql+pyodbc:///?odbc_connect=%s" % paramAQ)


#**** BBDD (Luis R) *****

###### DW ######

param_DW = urllib.parse.quote_plus("DRIVER={SQL Server};SERVER=10.34.71.202;UID=datawarehouse;PWD=datawarehouse;DATABASE=DM_Transaccional;TrustServerCertificate=yes")
conn_DW = sa.create_engine("mssql+pyodbc:///?odbc_connect=%s" % param_DW)


##### DEV #####

param_dev = urllib.parse.quote_plus("DRIVER={/opt/microsoft/msodbcsql18/lib64/libmsodbcsql-18.2.so.2.1};SERVER=10.34.71.227;UID=eymetric;PWD=Eym3tr1c.2022;DATABASE=DM_Transaccional_2022;TrustServerCertificate=yes")
conn_dev = sa.create_engine("mssql+pyodbc:///?odbc_connect=%s" % param_dev)

##### AQUILES #####

param_AQ = urllib.parse.quote_plus("DRIVER={SQL Server};DATABASE=DCCPProcurement;SERVER=10.34.71.146;UID=datawarehouse;PWD=datawarehouse;DATABASE=DCCPProcurement;TrustServerCertificate=yes")
conn_AQ = sa.create_engine("mssql+pyodbc:///?odbc_connect=%s" % param_AQ)

### connections with pyodbc ###
conn_AQ_pyodbc=pyodbc.connect('DRIVER={SQL Server};DATABASE=DCCPProcurement;SERVER=10.34.71.146;UID=datawarehouse;PWD=datawarehouse;DATABASE=DCCPProcurement;TrustServerCertificate=yes')#;Encrypt=yes')#;TrustServerCertificate=yes')#;TrustServerCertificate=yes')
#conn_AQ_pyodbc = pyodbc.connect('DRIVER={SQL Server};SERVER=10.34.71.146;DATABASE=DCCPProcurement;UID=datawarehouse;PWD=datawarehouse')
cursor=conn_AQ_pyodbc.cursor()
#*************************

# =============================================================================
# Creando parametros y funciones para queries 
# =============================================================================


#url = "C:/Users/luis.recabarren/Documents/proyecto_statistics_tool/output/"
#url='/mnt/c/Users/luis.recabarren/Documents/proyecto_statistics_tool/output/'
#url = 'C:/Users/USER/OneDrive - DCCP/Documents/comunicaciones dpto/Reporte_cifras_MP_automatizado'
url = 'C:/o/OneDrive - DCCP/Escritorio/Dpt Comunicaciones/Reportes Automatizados'

AnoReg = 2023

#Funcion para reducir un Anio
def Anomenos1(Ano):
    Menos = Ano - 1
    return Menos

#Funcion para poner puntos decimales
def itanum(x):
    return format(x,',d').replace(",",".")    


AnoRegM1 = Anomenos1(AnoReg)

#Si es semestral, agregar meses correspondientes
MesInferior = 1
MesSuperior = 8

#Agregar que semestre es:
Sem = "primer"
tri='primer' #'segundo' ,etc

# =============================================================================
# Fin Inputs
# =============================================================================

########################################
### Query con datos a nivel regional ###
########################################

FechaQ = pd.Series([AnoReg,AnoRegM1, MesInferior,MesSuperior])




#cursor.execute("EXEC "+upside_sproc)
##Se utiliza el wildcard ? para agregar los parámetros
QueryReg = pd.read_sql(con = conn_DW,  sql = '''
    ;with temp as (
	SELECT TPO.YEAR 'Ano'
         , tpo.Month as 'Mes'
        , LOC.Region
       , (CASE oc1.porisintegrated WHEN 3 THEN 'CAg'
                else (case  OC1.IDProcedenciaOC
                                WHEN 703 THEN 'Convenio Marco'
                                WHEN 701 THEN 'Licitación Pública'
                                WHEN 1401 THEN 'Licitación Pública'
                                WHEN 702 THEN 'Licitación Privada'
                                ELSE 'Trato Directo' END)END) 'ProcedenciaOC'
       , I.NombreInstitucion
       ,case when a.idTamano is not null then cast(a.idTamano as nvarchar)
				 else isnull(b.tamanonombre,'SinDato') end tamano0
       ,  COUNT(DISTINCT P.RUTSucursal)      'CantProveedores'
       , COUNT(DISTINCT OC1.CodigoOC) 'CantOC'    
       , SUM(OC1.MontoUSD+OC1.ImpuestoUSD) 'MONTOUSD'
       , SUM(OC1.MontoCLP+OC1.ImpuestoCLP) 'MONTOCLP'
       , SUM(OC1.MontoCLF+OC1.ImpuestoCLF) 'MONTOCLF'

 

    FROM DM_Transaccional..THOrdenesCompra  AS OC1
    inner JOIN DM_Transaccional..DimTiempo AS TPO                     ON OC1.IDFechaEnvioOC=TPO.DateKey
    inner JOIN DM_Transaccional..DimProcedenciaOC AS PRO   ON PRO.IDProcedenciaOC = OC1.IDProcedenciaOC
    left JOIN DM_Transaccional..DimComprador AS C        ON OC1.IDUnidaddeCompra = C.IDUnidaddeCompra
    left JOIN DM_Transaccional..DimProveedor AS P        ON P.IDSucursal=OC1.IDSucursal
	left join [DM_Transaccional].[dbo].[THTamanoProveedor] a on a.entcode=p.entCode and AñoTributario=2021
    left join Estudios.dbo.TamanoProveedorNuevos20230802 b on p.entcode=b.entCode
    LEFT JOIN  DM_Transaccional..DimTamanoProveedor AS TP ON TP.IdTamano=a.IdTamano
    left JOIN DM_Transaccional..DimInstitucion AS I      ON C.entCode = I.entCode
    left JOIN DM_Transaccional..DimSector AS S           ON S.IdSector = I.IdSector
    left JOIN DM_Transaccional..DimLocalidad as loc      ON C.IDLocalidadUnidaddeCompra =  LOC.IDLocalidad

    --Dado que existe comparacion anual, debe tomarse con año 1 y año -1
    WHERE   TPO.YEAR in(2023,2022)
            AND TPO.MONTH >= 1 AND TPO.MONTH <= 8
    GROUP BY TPO.YEAR,TPO.MONTH, 
	case when a.idTamano is not null then cast(a.idTamano as nvarchar)
				 else isnull(b.tamanonombre,'SinDato') end ,
				 I.NombreInstitucion,  LOC.Region,
           (CASE OC1.porisintegrated WHEN 3 THEN 'CAg'
                else (case  OC1.IDProcedenciaOC
                                WHEN 703 THEN 'Convenio Marco'
                                WHEN 701 THEN 'Licitación Pública'
                                WHEN 1401 THEN 'Licitación Pública'
                                WHEN 702 THEN 'Licitación Privada'
                                ELSE 'Trato Directo' END)END)
    )
	Select 		Ano

        , Mes
        , Region
       ,  ProcedenciaOC
       , NombreInstitucion
	   ,case  when tamano0 in ('PYME','Microempresa','2','3','4') then 'MiPyme'
				when tamano0 in ('Grande','1') then 'Grande'
				when tamano0  in ('5','Extranjera','SinDato') then 'SinDato' 
					end Tamano,
			sum(MONTOUSD) MONTOUSD,
			sum(MONTOCLF) MONTOCLF,
			sum(MONTOCLP) MONTOCLP,
			sum(CantOC) CantOC,
			sum(CantProveedores)      'CantProveedores'
	from temp
	group by 	Ano   
        , Mes
        , Region
       ,  ProcedenciaOC
       , NombreInstitucion
	  , case  when tamano0 in ('PYME','Microempresa','2','3','4') then 'MiPyme'
				when tamano0 in ('Grande','1') then 'Grande'
				when tamano0  in ('5','Extranjera','SinDato') then 'SinDato' 
					end
    ''')

#sys.exit()


#####################################
### Query top 3 rubros por region ###
#####################################

##Parametros de fecha Query rubro
FechaQRubro = FechaQ.drop(FechaQ.index[1])


QueryRegRubro = pd.read_sql(con = conn_DW,  sql = 
''' 
select * 
 from( select
LOC.Region
,rubro.RubroN1
,SUM(OCL.MontoUSD) 'MONTOUSD'
, SUM(OCL.MontoCLP) 'MONTOCLP'
, SUM(OCL.MontoCLF) 'MONTOCLF'
, Rank() over (Partition BY LOC.Region ORDER BY  SUM(OCL.MontoUSD) DESC ) AS Rank
	FROM  
	DM_Transaccional..THOrdenesCompra AS OC     
	LEFT JOIN DM_Transaccional..DimTiempo AS TPO                     ON OC.IDFechaEnvioOC=TPO.DateKey 
	LEFT JOIN DM_Transaccional..DimProcedenciaOC AS PRO   ON PRO.IDProcedenciaOC = OC.IDProcedenciaOC
	INNER JOIN DM_Transaccional..DimComprador AS C        ON OC.IDUnidaddeCompra = C.IDUnidaddeCompra
	INNER JOIN DM_Transaccional..DimProveedor AS P        ON P.IDSucursal=OC.IDSucursal
	--LEFT JOIN  DM_Transaccional..DimTamanoProveedor AS TP ON TP.IdTamano=P.IdTamano
	INNER JOIN DM_Transaccional..DimInstitucion AS I      ON C.entCode = I.entCode   
	INNER JOIN DM_Transaccional..DimSector AS S           ON S.IdSector = I.IdSector
	INNER JOIN DM_Transaccional..DimLocalidad as loc	  ON C.IDLocalidadUnidaddeCompra =  LOC.IDLocalidad
	LEFT JOIN  DM_Transaccional..THOrdenesCompraLinea OCL ON OCL.porID=OC.porID
	LEFT JOIN DM_Transaccional..DimProducto produc ON produc.IDProducto=OCL.IDProducto 
	LEFT JOIN DM_Transaccional..DimRubro rubro on rubro.IdRubro=produc.IdRubro
	--Dado que existe comparacion anual, debe tomarse con año 1 y año -1
	WHERE   TPO.YEAR = 2023
    	AND  TPO.month between 1 and 8
	GROUP BY LOC.Region,rubro.RubroN1
			)rs WHERE Rank <= 3 ORDER BY Region ASC, Rank ASC
                                               ''') 


#sys.exit()

##Cambio formato df
##QueryRegRubro['Region'] = str(QueryRegRubro['Region'])
 
##############################################################
### Query con datos a nivel nacional. Solo no tiene Region ###
##############################################################

print('Datos a nivel nacional. Inputs:\n')
print(FechaQ)
print()

QueryTotal = pd.read_sql(con = conn_DW, sql = '''
;with temp as (
    SELECT TPO.YEAR 'Ano' --TPO.MONTHNAME 'Mes'
    , tpo.Month as 'Mes'
    , (CASE oc1.porisintegrated WHEN 3 THEN 'CAg'
            else (case  OC1.IDProcedenciaOC
                            WHEN 703 THEN 'Convenio Marco'
                                WHEN 701 THEN 'Licitación Pública'
                                WHEN 1401 THEN 'Licitación Pública'
                                WHEN 702 THEN 'Licitación Privada'
                                ELSE 'Trato Directo' END)END) 'ProcedenciaOC'
	   ,I.NombreInstitucion
       ,case when a.idTamano is not null then cast(a.idTamano as nvarchar)
                    else isnull(b.tamanonombre,'SinDato') end tamano0
       , COUNT(DISTINCT P.RUTSucursal)      'CantProveedores'
       , COUNT(DISTINCT OC1.CodigoOC) 'CantOC'    
       , SUM(OC1.MontoUSD+OC1.ImpuestoUSD) 'MONTOUSD'
       , SUM(OC1.MontoCLP+OC1.ImpuestoCLP) 'MONTOCLP'
       , SUM(OC1.MontoCLF+OC1.ImpuestoCLF) 'MONTOCLF'

    FROM DM_Transaccional..THOrdenesCompra  AS OC1
    inner JOIN DM_Transaccional..DimTiempo AS TPO                     ON OC1.IDFechaEnvioOC=TPO.DateKey
    inner JOIN DM_Transaccional..DimProcedenciaOC AS PRO   ON PRO.IDProcedenciaOC = OC1.IDProcedenciaOC
    left JOIN DM_Transaccional..DimComprador AS C        ON OC1.IDUnidaddeCompra = C.IDUnidaddeCompra
    left JOIN DM_Transaccional..DimProveedor AS P        ON P.IDSucursal=OC1.IDSucursal
    left join [DM_Transaccional].[dbo].[THTamanoProveedor] a on a.entcode=p.entCode and AñoTributario=2021
    left join Estudios.dbo.TamanoProveedorNuevos20230802 b on p.entcode=b.entCode
    LEFT JOIN  DM_Transaccional..DimTamanoProveedor AS TP ON TP.IdTamano=a.IdTamano
    left JOIN DM_Transaccional..DimInstitucion AS I      ON C.entCode = I.entCode
    left JOIN DM_Transaccional..DimSector AS S           ON S.IdSector = I.IdSector
    left JOIN DM_Transaccional..DimLocalidad as loc      ON C.IDLocalidadUnidaddeCompra =  LOC.IDLocalidad
    
	--Dado que existe comparacion anual, debe tomarse con año 1 y año -1
    WHERE   TPO.YEAR in(2023,2022)
            AND TPO.MONTH >= 1 AND TPO.MONTH <= 8
    GROUP BY TPO.YEAR,TPO.MONTH, TP.Tamano,
	case when a.idTamano is not null then cast(a.idTamano as nvarchar)
				 else isnull(b.tamanonombre,'SinDato') end ,
				 I.NombreInstitucion,
           (CASE oc1.porisintegrated WHEN 3 THEN 'CAg'
                else (case  OC1.IDProcedenciaOC
                                WHEN 703 THEN 'Convenio Marco'
                                WHEN 701 THEN 'Licitación Pública'
                                WHEN 1401 THEN 'Licitación Pública'
                                WHEN 702 THEN 'Licitación Privada'
                                ELSE 'Trato Directo' END)END)
   )
   Select       Ano

        , Mes
       ,  ProcedenciaOC
       , NombreInstitucion
       ,case  when tamano0 in ('PYME','Microempresa','2','3','4') then 'MiPyme'
                when tamano0 in ('Grande','1') then 'Grande'
                when tamano0  in ('5','Extranjera','SinDato') then 'SinDato'
                    end Tamano,
            sum(MONTOUSD) MONTOUSD,
            sum(MONTOCLF) MONTOCLF,
            sum(MONTOCLP) MONTOCLP,
            sum(CantOC) CantOC,
            sum(CantProveedores)      'CantProveedores'
    from temp
    group by    Ano
        , Mes
       ,  ProcedenciaOC
       , NombreInstitucion
      , case  when tamano0 in ('PYME','Microempresa','2','3','4') then 'MiPyme'
                when tamano0 in ('Grande','1') then 'Grande'
                when tamano0  in ('5','Extranjera','SinDato') then 'SinDato'
                    end
        ''')

print(QueryTotal)
print()
print('--------------------------------------------------------------------------------------------------------------------------------------------------------------------')
print()
#Se requiere una query aparte para calcular el total de proveedores. Se quita mes y procedencia

########################
### Total  Proveedor ###
########################

FechaTotProv =  FechaQ.drop(FechaQ.index[1])
print('Tamaño proveedor (nacional). Inputs:\n')
print(FechaTotProv)
print()

QueryTotalProv = pd.read_sql(con = conn_DW,  sql = 
        '''SELECT       
				case 
					when a.idTamano is not null then 
						case a.idTamano when 1 then 'Grande'
						else 'Mipyme' end
				else 
					case isnull(b.Tamano,5) 
						when 1 then 'Grande'
						else 'Mipyme' end end Tamano

			,SUM(OC.MontoUSD+OC.ImpuestoUSD) 'MONTOUSD'
		   , SUM(OC.MontoCLF+OC.ImpuestoCLF) 'MONTOCLF'
		   ,SUM(OC.MontoCLP+OC.ImpuestoCLP) 'MONTOCLP'
			  ,count(distinct oc.porid) CantOC
               ,COUNT(DISTINCT P.RUTSucursal) 'CantProveedores'
			--  into #aux1
     FROM  
    DM_Transaccional..THOrdenesCompra AS OC     
    inner JOIN DM_Transaccional..DimTiempo AS TPO                     ON OC.IDFechaEnvioOC=TPO.DateKey 
	inner join [DM_Transaccional].[dbo].[dimproveedor] p on p.orgCode=oc.IDSucursal
    left join   [DM_Transaccional].[dbo].[THTamanoProveedor] a on a.entcode=p.entCode and AñoTributario=2021
    left join Estudios.dbo.TamanoProveedorNuevos20230809 b on p.entcode=b.entCode
    WHERE   TPO.YEAR in (2023) and TPO.MONTH >= 1 AND TPO.MONTH <= 8
	group by case 
					when a.idTamano is not null then 
						case a.idTamano when 1 then 'Grande'
						else 'Mipyme' end
				else 
					case isnull(b.Tamano,5) 
						when 1 then 'Grande'
						else 'Mipyme' end end    
                             ''')
print(QueryTotalProv)
print()
print('--------------------------------------------------------------------------------------------------------------------------------------------------------------------')
print()
#sys.exit()

##################################
### Query top 10 OC por region ###
##################################

FechaQTop10 = FechaQ.drop(FechaQ.index[1])
print('Top 10 OC por region','Inputs:\n')
print(FechaQTop10)
print()

QTop10 = pd.read_sql(con = conn_DW,  sql = ''' 
SELECT *
    FROM (
        SELECT 
                case 
					when a.idTamano is not null then 
						case a.idTamano when 1 then 'Grande'
						else 'Mipyme' end
				else 
					case isnull(b.Tamano,5) 
						when 1 then 'Grande'
						else 'Mipyme' end end Tamano,
    DM_Transaccional.dbo.DimInstitucion.NombreInstitucion,
    DM_Transaccional.dbo.DimComprador.NombreUnidaddeCompra,
    DM_Transaccional.dbo.DimLocalidad.Region,
    DM_Transaccional.dbo.DimProveedor.NombreSucursal,
    --DM_Transaccional.dbo.DimTamanoProveedor.Tamano,
    DM_Transaccional.dbo.THOrdenesCompra.CodigoOC,
    DM_Transaccional.dbo.THOrdenesCompra.NombreOC,
    DM_Transaccional.dbo.THOportunidadesNegocio.NombreAdq as NombreLic,
    ISNULL(DM_Transaccional.dbo.THOportunidadesNegocio.NombreAdq, DM_Transaccional.dbo.THOrdenesCompra.NombreOC) as MotivoCompra,
    round(DM_Transaccional.dbo.THOrdenesCompra.MontoUSD + DM_Transaccional.dbo.THOrdenesCompra.ImpuestoUSD,0) USD_BRUTO,
    DM_Transaccional.dbo.THOrdenesCompra.MontoCLP + DM_Transaccional.dbo.THOrdenesCompra.ImpuestoCLP PESOS_BRUTO, Rank()
          over (Partition BY DM_Transaccional.dbo.DimLocalidad.Region
                ORDER BY DM_Transaccional.dbo.THOrdenesCompra.MontoUSD DESC ) AS Rank
                FROM DM_Transaccional.dbo.DimInstitucion INNER JOIN
             DM_Transaccional.dbo.DimComprador ON DM_Transaccional.dbo.DimInstitucion.entCode = DM_Transaccional.dbo.DimComprador.entCode INNER JOIN
             DM_Transaccional.dbo.DimLocalidad ON DM_Transaccional.dbo.DimComprador.IDLocalidadUnidaddeCompra = DM_Transaccional.dbo.DimLocalidad.IDLocalidad INNER JOIN
             DM_Transaccional.dbo.THOrdenesCompra ON DM_Transaccional.dbo.DimComprador.IDUnidaddeCompra = DM_Transaccional.dbo.THOrdenesCompra.IDUnidaddeCompra INNER JOIN
             DM_Transaccional.dbo.DimProveedor ON DM_Transaccional.dbo.DimProveedor.IDSucursal = DM_Transaccional.dbo.THOrdenesCompra.IDSucursal INNER JOIN
             --DM_Transaccional.dbo.DimTamanoProveedor ON DM_Transaccional.dbo.DimTamanoProveedor.IdTamano = DM_Transaccional.dbo.DimProveedor.IdTamano INNER JOIN
             DM_Transaccional.dbo.DimTiempo ON DM_Transaccional.dbo.DimTiempo.DateKey =  DM_Transaccional.dbo.THOrdenesCompra.IDFechaEnvioOC LEFT JOIN
             DM_Transaccional.dbo.THOportunidadesNegocio ON DM_Transaccional.dbo.THOportunidadesNegocio.rbhCode = DM_Transaccional.dbo.THOrdenesCompra.rbhCode
             left join   [DM_Transaccional].[dbo].[THTamanoProveedor] a on a.entcode=[DM_Transaccional].[dbo].[dimproveedor].entCode and AñoTributario=2021
             left join Estudios.dbo.TamanoProveedorNuevos20230809 b on [DM_Transaccional].[dbo].[dimproveedor].entcode=b.entCode
          where DM_Transaccional.dbo.dimtiempo.Year = 2023
             and DM_Transaccional.dbo.dimtiempo.month between 1 and 8
        ) rs WHERE Rank <= 3 ORDER BY Region ASC, Rank ASC
                ''')


print(QTop10)
print()
print('--------------------------------------------------------------------------------------------------------------------------------------------------------------------')
print()
#sys.exit()





###########################################################################
### Query para obtener número total de proveedores que participan en MP ###
###########################################################################

##Calculando fecha para query dinamica
FechaProv =  FechaQ.drop(FechaQ.index[1])
##Proveedores solo transando
print('Proveedores transando en MP.','Inputs:\n')
print(FechaProv)
print()
#vec=['2023','1','6'] ### vector de parametros para usar con con=conn_AQ_pyodbc

QProvTrans = pd.read_sql(con = conn_AQ,  sql = '''
SELECT DISTINCT C.entCode AS Transan
FROM         
DCCPProcurement..prcPOHeader A with(nolock)INNER JOIN
DCCPPlatform..gblOrganization B with(nolock) ON A.porSellerOrganization = B.orgCode INNER JOIN
DCCPPlatform..gblEnterprise C with(nolock) ON B.orgEnterprise = C.entCode
WHERE     (A.porBuyerStatus IN (4, 5, 6, 7, 12)) 
AND year(A.porSendDate) =  2023 
and (month(A.porSendDate)>=1 AND month(A.porSendDate)<= 8 )
''')
print(QProvTrans)
print()
print('--------------------------------------------------------------------------------------------------------------------------------------------------------------------')
print()
#sys.exit()

#FechaProv2=pd.concat([FechaProv],ignore_index=True) #,FechaProv,FechaProv,FechaProv], ignore_index=True) # For conn_AQ connection using
#print()
#
#vec2=[] # vector de parametros para conn_AQ_pyodbc
#for k in range(1):
#    vec2.append(vec[0])
#    vec2.append(vec[1])
#    vec2.append(vec[2])
#
#
#QProv = pd.read_sql(con = conn_AQ, params= FechaProv2, sql = '''
#
#    /*Proveedores involucrados en una OC*/
#    SELECT DISTINCT C.entCode AS Transan
#    FROM         
#    DCCPProcurement..prcPOHeader A with(nolock)INNER JOIN
#    DCCPPlatform..gblOrganization B with(nolock) ON A.porSellerOrganization = B.orgCode INNER JOIN
#    DCCPPlatform..gblEnterprise C with(nolock) ON B.orgEnterprise = C.entCode
#    WHERE     (A.porBuyerStatus IN (4, 5, 6, 7, 12)) 
#    AND year(A.porSendDate) =  ? 
#    and ( month(A.porSendDate)>= ? AND month(A.porSendDate)<= ? )
#    
#    UNION 
#    
#    /*Proveedores que han participado emitiendo una oferta*/
#    
#    SELECT DISTINCT C.orgEnterprise as Transan
#    FROM         
#    DCCPProcurement..prcBIDQuote A with(nolock)INNER JOIN
#    DCCPProcurement..prcRFBHeader B with(nolock) ON A.bidRFBCode = B.rbhCode INNER JOIN
#    DCCPPlatform..gblOrganization C with(nolock) ON A.bidOrganization = C.orgCode
#    WHERE     (A.bidDocumentStatus IN (3, 4, 5)) 
#    AND year(A.bidEconomicIssueDate) = ? 
#    and (month(A.bidEconomicIssueDate)>= ? AND month(A.bidEconomicIssueDate)<= ? )
#
#    UNION
#    
#    /*Proveedores a los que se les ha solicitado una cotización*/
#    
#    SELECT  DISTINCT (B.orgEnterprise) AS Transan     
#      FROM [DCCPProcurement].[dbo].[prcPOCotizacion] A
#      INNER JOIN DCCPPlatform..gblOrganization B ON
#      A.proveedorRut=B.orgTaxID INNER JOIN
#      DCCPProcurement..prcPOHeader C ON
#      A.porId = C.porID
#      WHERE year(C.porSendDate) = ? 
#    and (month(C.porSendDate)>= ? AND month(C.porSendDate)<= ?)
#    
#    UNION
#    /*MÁS LOS RUT NO REGISTRADOS EN MERCADO PÚBLICO Y QUE SON DEL PERÍODO DE REFERENCIA*/
#    
#    SELECT  
#    DISTINCT A.proveedorRut   
#      FROM [DCCPProcurement].[dbo].[prcPOCotizacion] A INNER JOIN
#       DCCPProcurement..prcPOHeader C ON
#      A.porId = C.porID
#      WHERE A.proveedorRut NOT IN (
#    SELECT  
#    DISTINCT A.proveedorRut     
#      FROM [DCCPProcurement].[dbo].[prcPOCotizacion] A
#      INNER  JOIN DCCPPlatform..gblOrganization B ON
#      A.proveedorRut=B.orgTaxID) AND 
#      year(C.porSendDate) = ? and 
#      (month(C.porSendDate)>= ? AND month(C.porSendDate)<= ?)''')
#
##Query total proveedores que participaron (No solo tienen OC asociada) Pedir a Javier
#print(QProv)
#print()
#print('--------------------------------------------------------------------------------------------------------------------------------------------------------------------')
#print()
##sys.exit()


####  ESTA ES EL OUTPUT DE LA QUERY ANTERIOR PERO YA EJECUTADA EN SQL SERVER, PORQUE LA ANTERIOR DEMORA MUCHO ####
QProv=pd.read_sql(con=conn_AQ, sql='''
    SELECT *
    FROM estudios.dbo.ProveedoresTransando20230811
    ''')
print(QProv)
print()
print('--------------------------------------------------------------------------------------------------------------------------------------------------------------------')
print()
#sys.exit()



# =============================================================================
# Metricas Nacionales
# =============================================================================

#Se utiliza la QueryTotal para las operaciones
#Se debe filtrar por el ano actual

QueryTotalA = QueryTotal.loc[QueryTotal['Ano'] == AnoReg]

#Total transado nacional en USD, CLP , CLF
NacTotTransUSDA = QueryTotalA['MONTOUSD'].sum()
NacTotTransCLPA = QueryTotalA['MONTOCLP'].sum()
NacTotTransCLFA = QueryTotalA['MONTOCLF'].sum ()

#Total OC
NacTotOCA = QueryTotalA['CantOC'].sum()


#Total proveedores con OC
NacProvTrans = QueryTotalProv['CantProveedores'].sum()


#Proporcion numero prov pyme vs total
NacTotProvA = QueryTotalProv['CantProveedores'].sum()

#Porcentaje de Pyme vs total proveedores
NacFiltPymeA = QueryTotalProv.loc[QueryTotalProv['Tamano'].isin(['MiPyme'])]
#Cuento las filas de la columna Ano como referencia, puede ser cualquier otra
NacTotPymeA = NacFiltPymeA['CantProveedores'].sum()
NacVarProvPymeA = (NacTotPymeA/NacTotProvA)*100

#Total proveedores transando 

#NacProvPart = QProv['Transan'].count()
NacProvPart=len(QProv)

#Total monto Convenio Marco
NacFiltCM = QueryTotalA[QueryTotalA['ProcedenciaOC'] == 'Convenio Marco']
NacFiltCMUSD = NacFiltCM['MONTOUSD'].sum()
NacFiltCMCLP = NacFiltCM['MONTOCLP'].sum()
             

NacFiltPymeTransA = QueryTotalA.loc[QueryTotalA['Tamano'].isin(['MiPyme'])]                                  
#Transado por Pyme
NacTransPymeUSDA =  NacFiltPymeTransA['MONTOUSD'].sum()
NacTransPymeCLFA =  NacFiltPymeTransA['MONTOCLF'].sum()
NacTransPymeCLPA =  NacFiltPymeTransA['MONTOCLP'].sum()

#Var transado pyme total
NacVarTransPymeA = (NacTransPymeCLFA/NacTotTransCLFA)*100

#Top 3 participacion Mipymes (Usa qery de regiones)
#Arreglo de regiones
ListReg = list(QueryReg['Region'].unique())
#DF vacio
NacPymePart = pd.DataFrame(columns = ['Region','PYME','Grande','Total'])

#Itero por region calculando porcentage de PYME en los montos totales
for i in ListReg:
#    print(i)
    RegFiltIter = QueryReg.loc[(QueryReg['Region'] == i) & (QueryReg['Ano'] == AnoReg)]
    RegFiltPymeA = RegFiltIter.loc[RegFiltIter['Tamano'].isin(['MiPyme'])]
    
    RegTotTransCLFA = RegFiltIter['MONTOCLF'].sum()
    RegTotTransPymeCLFA = RegFiltPymeA['MONTOCLF'].sum()
    VarTransPymeA = (RegTotTransPymeCLFA/RegTotTransCLFA)*100
    VarGrande = 100-VarTransPymeA
    NuevaFila = pd.DataFrame({'Region':'Atacama', 'PYME':VarTransPymeA, 'Grande': VarGrande, 'Total': 100},index=[0])
    NacPymePart= pd.concat([NuevaFila,NacPymePart.loc[:]]).reset_index(drop=True)
    #NacPymePart = NacPymePart.concat({'Region':i, 'PYME':VarTransPymeA, 'Grande': VarGrande, 'Total': 100} , ignore_index=True)



NacPymePart = NacPymePart.sort_values('PYME', ascending = False) 
   
#Generar top 3 pyme
NacTop3Pyme = NacPymePart.head(3)

# =============================================================================
# Grafico pyme en Seaborn
# =============================================================================

#DF a utilizar: NacPymePart
#Cambiando columna a index para grafico
sns.set(style="whitegrid")

# Initialize the matplotlib figure
f, ax = plt.subplots(figsize=(6, 6))
#ax.invert_xaxis()

# Plot the total crashes
sns.set_color_codes("deep")
#Asignar a variable primer grafico
GrafPyme =sns.barplot(x='Total', y="Region", data=NacPymePart,
            label="Grande", color="b")

# Plot the crashes where alcohol was involved
sns.set_color_codes("muted")
splot=sns.barplot(x="PYME", y="Region", data=NacPymePart,
            label="MiPyme", color="r")
#splot.set_title("Montos 2021 MiPyme sobre total transado por región", fontweight="bold"
#                                   ,fontsize = 12)



#anotar cantidad de porcentaje por barra
for p in itertools.islice(splot.patches,16,32):
    splot.annotate(format(p.get_width(), '.1f')+'%', # formato 
                   (p.get_width() - 10, p.get_y()-0.1), # coordenada donde imprimir la cifra
                   ha = 'center', va = 'center', 
                   size=15,
                   xytext = (0, -12), 
                   textcoords = 'offset points')



# Add a legend and informative axis label
ax.legend(ncol=2, loc="lower right", frameon=True,bbox_to_anchor=(-0.02, -0.1))
ax.set(xlim=(0, 100), ylabel="",
       xlabel="Porcentaje")
sns.despine(left=True, bottom=True)
#Ajustando espacio para que regiones no sean cortadas
plt.subplots_adjust(left=0.4, right =0.9 )




fig = GrafPyme.get_figure()
fig.savefig(url+'GrafPyme.jpg') 

#Pymechart.savefig(url+"output.png")
#Guardar chart
    
# =============================================================================
# Generar Iterador  con metricas Regionales
# =============================================================================
                                  
#Iterando para que r tome el valor de todas las regiones disponibles
for r in ListReg:

    Reg = r
    #Reg = 'Antofagasta'
                                                      
    #Convencion: A = Año presente , B = Año pasado         
    RegFiltA = QueryReg.loc[(QueryReg['Region'] == Reg) & (QueryReg['Ano'] == AnoReg)]
    
    #Total transado por region en USD, CLP , CLF
    RegTotTransUSDA = RegFiltA['MONTOUSD'].sum()
    RegTotTransCLPA = RegFiltA['MONTOCLP'].sum()
    RegTotTransCLFA = RegFiltA['MONTOCLF'].sum()
    
    #Calculo de monto en CLF para calcular variacion con año pasado
    RegFiltB = QueryReg.loc[(QueryReg['Region'] == Reg) & (QueryReg['Ano'] == AnoRegM1)]
    RegTotTransCLFB = RegFiltB['MONTOCLF'].sum()
    
    #Calculo de variacion anual
    VarRegCLF = (RegTotTransCLFA - RegTotTransCLFB)/ RegTotTransCLFB *100
    
    #Total OC region
    RegTotOCA = RegFiltA['CantOC'].sum()
    
    #Calculando datos pyme
    RegFiltPymeA = RegFiltA.loc[RegFiltA['Tamano'].isin(['MiPyme'])]
    RegTotTransPymeCLFA = RegFiltPymeA['MONTOCLF'].sum()
    RegTotTransPymeCLPA = RegFiltPymeA['MONTOCLP'].sum()
    #Variacion transado pyme vs total transado
    VarTransPymeA = (RegTotTransPymeCLFA/RegTotTransCLFA)*100
    
    #Ranking top 3 instituciones con mayor monto transado
    RankFiltTransA = RegFiltA.groupby(by = 'NombreInstitucion')['MONTOCLP'].sum().sort_values(ascending = False).head(3)
    RankFiltTransA = RankFiltTransA.to_frame()
    RankFiltTransA= RankFiltTransA.reset_index()
    RankFiltTransA['NombreInstitucion'] = RankFiltTransA['NombreInstitucion'].str.upper()
    
    #Ranking top 3 Rubros con mayor monto transado
    RankRubroReg =  QueryRegRubro.loc[QueryRegRubro['Region'] == Reg]
    
    #Ranking top 3 OC de region para word
    RegRankOC = QTop10.loc[QTop10['Region'] == Reg]
    RegRankOC = RegRankOC[['NombreInstitucion','CodigoOC','MotivoCompra','NombreSucursal','USD_BRUTO']]
    
    #Version para word
    
    RegRankOCWord = RegRankOC.rename({'NombreInstitucion': 'Institución', 'CodigoOC': 'Código orden de compra'
                                      ,'MotivoCompra':'Motivo de la compra'
                                      ,'NombreSucursal': 'Proveedor','USD_BRUTO': 'Monto comprometido en US$'}, axis=1)
    
        
    #Cambiando formato de Dolares desde float a str
    RegRankOCWord['Monto comprometido en US$']  =  RegRankOCWord['Monto comprometido en US$'].astype(int) 
    RegRankOCWord['Monto comprometido en US$'] = RegRankOCWord['Monto comprometido en US$'].apply(itanum) 
    RegRankOCWord['Monto comprometido en US$']  = RegRankOCWord['Monto comprometido en US$'].astype(str)
    
    
    
    
    
    # =============================================================================
    # Crear Documento word
    # =============================================================================
    
    
    #Crear objeto documento  ()
    #Crear sobre un template para que quede con los logos de chilecompra
    #Obs, no se pueded
    document = docx.Document('template.docx')
    
    #"Agregar espacio
    document.add_heading(" " , 0)
    #Agregar titulo. el 0 indica el tipo de titulo
    document.add_heading('Datos regionales '  + str(AnoReg) + " " + str(Reg)  , 0)
    #document.add_heading('Datos regionales ' + str(AnoReg) + " " + str(Reg)  , 0)
    #document.add_heading('Datos regionales ' +Sem+ ' Trimestre ' + str(AnoReg) + "   " + str(Reg)  , 0)
    
    #Agregar primer título en 
    document.add_heading('Cifras Regionales', 1)
    #Generar espacios
    document.add_paragraph(" ")
    #Agregar texto en negra. Solo se agregan fuentes a nivel de run
    parBold1 = document.add_paragraph()
    #Agrega espacio desúés de paragraph

    ### Semestre ###
    runBold1 =parBold1.add_run("US$ " + str(itanum(int(RegTotTransUSDA))) + " se transaron en Mercado Público entre enero y agosto de " 
                                      + str(AnoReg) + " en la región de " +str(Reg)+".")
    
    ### Trimestre ###
    #runBold1 =parBold1.add_run("US$ " + str(itanum(int(RegTotTransUSDA))) + " se transaron en Mercado Público durante el " + Sem +  " trimestre de " 
    #                                  + str(AnoReg) + " en la región de " +str(Reg)+".")
    
    ### Ano ###
    #runBold1 =parBold1.add_run("US$ " + str(itanum(int(RegTotTransUSDA))) + " se transaron en Mercado Público durante el " 
    #                                  + str(AnoReg) + " en la región de " +str(Reg)+".")
    
    runBold1.bold = True
    parBold1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(" ")
    #parBold1.font.size = Pt(12)
    #Segundo paragraph solo en 
    parBold2 = document.add_paragraph()
    parBold2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    runBold2 = parBold2.add_run("Estas cifras consideran las diversas modalidades de compra existentes en el Estado, las cuales hoy ponen el énfasis en maximizar los ahorros en el fisco y hacer más ágiles los procesos para todos los ciudadanos.")
    runBold2.bold = True
    document.add_paragraph(" ")
    #Tercer paragraph
    
    ### Semestre ###
    par1Reg = document.add_paragraph("Entre enero y agosto de " + str(AnoReg) + " se transaron " + str(itanum(int(RegTotTransUSDA))) + " dólares equivalentes a "+ 
                                     str(itanum(int(RegTotTransCLPA))) + " pesos a través de Mercado Público en la Región de "+ str(Reg) +
                                     ". De ese total, las micro, pequeñas y medianas empresas se adjudicaron "+str(itanum(int(RegTotTransPymeCLPA)))+
                                     " pesos, es decir, un "+str(itanum(int(VarTransPymeA)))+"% del monto total.") 
    ### Ano ###
    #par1Reg = document.add_paragraph("Durante el " + str(AnoReg) + " se transaron " + str(itanum(int(RegTotTransUSDA))) + " dólares equivalentes a "+ 
    #                                 str(itanum(int(RegTotTransCLPA))) + " pesos a través de Mercado Público en la Región de "+ str(Reg) +
    #                                 ". De ese total, las micro, pequeñas y medianas empresas se adjudicaron "+str(itanum(int(RegTotTransPymeCLPA)))+
    #                                 " pesos, es decir, un "+str(itanum(int(VarTransPymeA)))+"% del monto total.") 
    
    ### Trimsestre ###
    #par1Reg = document.add_paragraph("Durante el " +Sem+ " trimestre de " + str(AnoReg) + " se transaron " + str(itanum(int(RegTotTransUSDA))) + " dólares equivalentes a "+ 
    #                                 str(itanum(int(RegTotTransCLPA))) + " pesos a través de Mercado Público en la Región de "+ str(Reg) +
    #                                 ". De ese total, las micro, pequeñas y medianas empresas se adjudicaron "+str(itanum(int(RegTotTransPymeCLPA)))+
    #                                 " pesos, es decir, un "+str(itanum(int(VarTransPymeA)))+"% del monto total.") 
    
    
    document.add_paragraph(" ")
    
    ### Semestre ###
    par2Reg = document.add_paragraph("Respecto al mismo periodo de "+ str(AnoRegM1)+
                                    ", se experimentó una variación de "+str(itanum(int(VarRegCLF)))+
                                    "% en términos reales, esto a partir de las "+str(itanum(int(RegTotOCA)))+
                                    " órdenes de compra emitidas por los diferentes organismos públicos.") 
   
    
    ### Ano ###
    #par2Reg = document.add_paragraph("Respecto al "+ str(AnoRegM1)+
    #                                 ", se experimentó una variación de "+str(itanum(int(VarRegCLF)))+
    #                                 "% en términos reales, esto a partir de las "+str(itanum(int(RegTotOCA)))+
    #                                 " órdenes de compra emitidas por los diferentes organismos públicos.")
    
    ### Trimestre ###
    #par2Reg = document.add_paragraph("Respecto al "+Sem+" trimestre del "+ str(AnoRegM1)+
    #                                 ", se experimentó una variación de "+str(itanum(int(VarRegCLF)))+
    #                                 "% en términos reales, esto a partir de las "+str(itanum(int(RegTotOCA)))+
    #                                 " órdenes de compra emitidas por los diferentes organismos públicos.") 
    
    document.add_paragraph(" ")
    ### Semestre ###
    par3Reg = document.add_paragraph("En la Región, en este "+
                                     " periodo, las instituciones con mayor participación fueron: "+RankFiltTransA.iloc[0]['NombreInstitucion']+
                                     " ("+ str(itanum(int(RankFiltTransA.iloc[0]['MONTOCLP'])))+ " pesos); "+RankFiltTransA.iloc[1]['NombreInstitucion']+
                                     " ("+ str(itanum(int(RankFiltTransA.iloc[1]['MONTOCLP'])))+ " pesos); " +RankFiltTransA.iloc[2]['NombreInstitucion']+
                                     " ("+ str(itanum(int(RankFiltTransA.iloc[2]['MONTOCLP'])))+
                                     " pesos). Por su parte, los rubros más solicitados fueron: "+RankRubroReg.iloc[0]['RubroN1']+
                                     " (US$ "+str(itanum(int(RankRubroReg.iloc[0]['MONTOUSD'])))+"); "+RankRubroReg.iloc[1]['RubroN1']+
                                     " (US$ "+str(itanum(int(RankRubroReg.iloc[1]['MONTOUSD'])))+"); "+RankRubroReg.iloc[2]['RubroN1']+
                                     " (US$ "+str(itanum(int(RankRubroReg.iloc[2]['MONTOUSD'])))+")."
                                     )
   
    
    ### Ano ###
    #par3Reg = document.add_paragraph("En la Región, las instituciones con mayor participación fueron: "+RankFiltTransA.iloc[0]['NombreInstitucion']+
    #                                 " ("+ str(itanum(int(RankFiltTransA.iloc[0]['MONTOCLP'])))+ " pesos); "+RankFiltTransA.iloc[1]['NombreInstitucion']+
    #                                 " ("+ str(itanum(int(RankFiltTransA.iloc[1]['MONTOCLP'])))+ " pesos); " +RankFiltTransA.iloc[2]['NombreInstitucion']+
    #                                 " ("+ str(itanum(int(RankFiltTransA.iloc[2]['MONTOCLP'])))+
    #                                 " pesos). Por su parte, los rubros más solicitados fueron: "+RankRubroReg.iloc[0]['RubroN1']+
    #                                 " (US$ "+str(itanum(int(RankRubroReg.iloc[0]['MONTOUSD'])))+"); "+RankRubroReg.iloc[1]['RubroN1']+
    #                                 " (US$ "+str(itanum(int(RankRubroReg.iloc[1]['MONTOUSD'])))+"); "+RankRubroReg.iloc[2]['RubroN1']+
    #                                 " (US$ "+str(itanum(int(RankRubroReg.iloc[2]['MONTOUSD'])))+")."
    #                                 )
    

    ### Trimestre ###
    #par3Reg = document.add_paragraph("En la Región, en este "+Sem+
    #                                " trimestre, las instituciones con mayor participación fueron: "+RankFiltTransA.iloc[0]['NombreInstitucion']+
    #                                " ("+ str(itanum(int(RankFiltTransA.iloc[0]['MONTOCLP'])))+ " pesos); "+RankFiltTransA.iloc[1]['NombreInstitucion']+
    #                               " ("+ str(itanum(int(RankFiltTransA.iloc[1]['MONTOCLP'])))+ " pesos); " +RankFiltTransA.iloc[2]['NombreInstitucion']+
    #                                 " ("+ str(itanum(int(RankFiltTransA.iloc[2]['MONTOCLP'])))+
    #                                 " pesos). Por su parte, los rubros más solicitados fueron: "+RankRubroReg.iloc[0]['RubroN1']+
    #                                 " (US$ "+str(itanum(int(RankRubroReg.iloc[0]['MONTOUSD'])))+"); "+RankRubroReg.iloc[1]['RubroN1']+
    #                                 " (US$ "+str(itanum(int(RankRubroReg.iloc[1]['MONTOUSD'])))+"); "+RankRubroReg.iloc[2]['RubroN1']+
    #                                 " (US$ "+str(itanum(int(RankRubroReg.iloc[2]['MONTOUSD'])))+")."
    #                                 )
    
    
    document.add_paragraph(" ")
    ### Semestre ###              
    par3Reg = document.add_paragraph("Las principales órdenes de compra, durante este periodo, fueron las siguentes:")     
    
    ### Ano ###
    #par3Reg = document.add_paragraph("Las principales órdenes de compra, durante el "+ str(AnoReg)+", fueron las siguentes:")     
    
    #xTrimestre
    #par3Reg = document.add_paragraph("Las principales órdenes de compra, durante el "+Sem+" trimestre, fueron las siguentes:")       
    
    document.add_paragraph(" ")
    #Tabla top OC por region
    #Indices parten en cero, igual que en python
    
    t = document.add_table(RegRankOCWord.shape[0]+1, RegRankOCWord.shape[1])
    
    # add the header rows.
    for j in range(RegRankOCWord.shape[-1]):
        t.cell(0,j).text = RegRankOCWord.columns[j]
    
    # add the rest of the data frame
    for i in range(RegRankOCWord.shape[0]):
        for j in range(RegRankOCWord.shape[-1]):
            t.cell(i+1,j).text = str(RegRankOCWord.values[i,j])    
        
    t.autofit = True
    t.style =  'EstiloDCCP'
    #document.add_page_break()
    #header contexto nacional
    document.add_heading(' ', 0)
    document.add_heading('Contexto nacional', 1)
    
    document.add_paragraph(" ")
    
    ### Semestre ###
    par1Nac = document.add_paragraph("A nivel nacional, entre enero y agosto del "+str(AnoReg)+
                                     " los montos totales transados por los organismos del Estado a través de la plataforma de ChileCompra, www.mercadopublico.cl, alcanzaron US$ "+
                                     str(itanum(int(NacTotTransUSDA)))+ " ("+str(itanum(int(NacTotTransCLPA)))+ " pesos).")
    
    document.add_paragraph(" ")
    par2Nac = document.add_paragraph("Las entidades públicas –ministerios, servicios, hospitales, municipios, universidades y FF.AA.- emitieron "+
                                    str(itanum(int(NacTotOCA)))+" órdenes de compra en la plataforma de compras públicas, a más de "+str(itanum(NacProvTrans))+" proveedores, el "+
                                     str(itanum(int(NacVarProvPymeA)))+"% de los cuales son micro y pequeñas empresas. El total de proveedores transando, es decir que emitieron ofertas, cotizaron y/o recibieron una orden de compra, fue de "+
                                      str(itanum(NacProvPart)) +" durante dicho periodo."+
                                     " Cabe recordar que las órdenes de compra en www.mercadopublico.cl, a partir de las cuales se calculan los montos, corresponden a compromisos y no a pagos de los organismos del Estado.")    
    
    
    ### Ano ###
    #par1Nac = document.add_paragraph("A nivel nacional, durante el "+str(AnoReg)+
    #                                 " los montos totales transados por los organismos del Estado a través de la plataforma de ChileCompra, www.mercadopublico.cl, alcanzaron US$ "+
    #                                 str(itanum(int(NacTotTransUSDA)))+ " ("+str(itanum(int(NacTotTransCLPA)))+ " pesos).")
    #
    #document.add_paragraph(" ")
    #par2Nac = document.add_paragraph("Las entidades públicas –ministerios, servicios, hospitales, municipios, universidades y FF.AA.- emitieron "+
    #                                 str(itanum(int(NacTotOCA)))+" órdenes de compra en la plataforma de compras públicas, a más de "+str(itanum(NacProvTrans))+" proveedores, el "+
    #                                 str(itanum(int(NacVarProvPymeA)))+"% de los cuales son micro y pequeñas empresas. El total de proveedores transando, es decir que emitieron ofertas, cotizaron y/o recibieron una orden de compra, fue de "+
    #                                  str(itanum(NacProvPart)) +" durante el "+str(AnoReg)+
    #                                ". Cabe recordar que las órdenes de compra en www.mercadopublico.cl, a partir de las cuales se calculan los montos, corresponden a compromisos y no a pagos de los organismos del Estado.")    
    
    ### Trimestre ###
    #par1Nac = document.add_paragraph("A nivel nacional, durante el "+Sem+" trimestre de "+str(AnoReg)+
    #                                " los montos totales transados por los organismos del Estado a través de la plataforma de ChileCompra, www.mercadopublico.cl, alcanzaron US$ "+
    #                                 str(itanum(int(NacTotTransUSDA)))+ " ("+str(itanum(int(NacTotTransCLPA)))+ " pesos).")
    
   # document.add_paragraph(" ")
   # par2Nac = document.add_paragraph("Las entidades públicas –ministerios, servicios, hospitales, municipios, universidades y FF.AA.- emitieron "+
   #                                  str(itanum(int(NacTotOCA)))+" órdenes de compra en la plataforma de compras públicas, a más de "+str(itanum(NacProvTrans))+" proveedores, el "+
   #                                  str(itanum(int(NacVarProvPymeA)))+"% de los cuales son micro y pequeñas empresas. El total de proveedores transando, es decir que emitieron ofertas, cotizaron y/o recibieron una orden de compra, fue de "+
   #                                   str(itanum(NacProvPart)) +" durante el "+Sem+
   #                                  " trimestre. Cabe recordar que las órdenes de compra en www.mercadopublico.cl, a partir de las cuales se calculan los montos, corresponden a compromisos y no a pagos de los organismos del Estado.")    
    
    
    
    document.add_paragraph(" ")        
    par3Nac = document.add_paragraph("A través de Convenios Marco –la tienda virtual más grande del país- los organismos públicos hicieron compras por más de US$ "+str(itanum(int(NacFiltCMUSD)))+
                                     " ("+str(itanum(int(NacFiltCMCLP)))+  " pesos)." )   
    
    document.add_paragraph(" ")
    par3Nac = document.add_paragraph()
    run3Nac = par3Nac.add_run("MiPymes lideran ventas con US$ "+str(itanum(int(NacTransPymeUSDA)))+" transados.")
    run3Nac.bold = True
    
    document.add_paragraph(" ")
    par4Nac = document.add_paragraph("A nivel nacional, la participación de las Mipymes (micro, pequeña y mediana) sobre el total de montos transados alcanzó en el periodo un "+
                                     str(itanum(int(NacVarTransPymeA)))+"%, cifra que corresponde a US$ "+str(itanum(int(NacTransPymeUSDA)))+" ("+str(itanum(int(NacTransPymeCLPA)))+
                                     " pesos). Esta es equivalente a más de tres veces la participación de este segmento en la economía nacional (15%). ")
    
    document.add_paragraph(" ")
    par5Nac = document.add_paragraph("En cuanto a la participación de las Mipymes sobre los montos transados de su región, destacan: " +
                                     NacPymePart.iloc[0]['Region']+", "+ NacPymePart.iloc[1]['Region']+", y "+
                                      NacPymePart.iloc[2]['Region']+ " con "+ str(itanum(int(NacPymePart.iloc[0]['PYME'])))+
                                      "%, "+ str(itanum(int(NacPymePart.iloc[1]['PYME'])))+"% y "+
                                      str(itanum(int(NacPymePart.iloc[2]['PYME'])))+"% respectivamente. ")
    
    document.add_paragraph(" ")
    
    #Agregar imagen de grafico
    prpyme = document.add_paragraph()
    runpyme = prpyme.add_run()
    #runpyme.add_picture(url+"GrafPyme.jpg",width=Inches(8.0), height=Inches(8.0)  )
    runpyme.add_picture(url+"GrafPyme.jpg")
    
    #Guardar documento
    urlDoc = url+'/output/'
    ### semestre ###
    document.save(urlDoc + str(Reg)+"_"+ Sem+ "-Sem-"+str(AnoReg)+".docx")
    
    ### anual ###
    #document.save(urlDoc + str(Reg)+"_"+str(AnoReg)+".docx")

    ### trimetre ###
    #document.save(urlDoc + str(Reg)+"_"+ Tri+ "-Trim-"+str(AnoReg)+".docx")



